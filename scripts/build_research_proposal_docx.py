from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "output" / "Research_Proposal_BI_Dashboard.docx"


TITLE = (
    "Design and Development of a Business Intelligence-Based Decision Support "
    "Dashboard for Strategic Decision-Making in Small and Medium Enterprises"
)


def set_run_font(run, size=12, bold=False, italic=False, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_format(paragraph, after=8, before=0, line_spacing=1.5, alignment=None):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing
    if alignment is not None:
        paragraph.alignment = alignment


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Page ")
    set_run_font(run, size=10)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def setup_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(8)

    for style_name, size, before, after in [
        ("Heading 1", 14, 14, 8),
        ("Heading 2", 13, 10, 6),
        ("Heading 3", 12, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.5

    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def add_para(doc, text="", bold_label=None, after=8, before=0, align=None):
    p = doc.add_paragraph()
    set_paragraph_format(p, after=after, before=before, alignment=align)
    if bold_label:
        label = p.add_run(bold_label)
        set_run_font(label, bold=True)
        if text:
            run = p.add_run(text)
            set_run_font(run)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_paragraph_format(p, after=4)
        run = p.add_run(item)
        set_run_font(run)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        set_paragraph_format(p, after=4)
        run = p.add_run(item)
        set_run_font(run)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size=14 if level == 1 else 13 if level == 2 else 12, bold=True)
    return p


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def format_table(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            cell.width = widths[idx]
            for paragraph in cell.paragraphs:
                set_paragraph_format(paragraph, after=0, line_spacing=1.2)
                for run in paragraph.runs:
                    set_run_font(run, size=10)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_timeline_table(doc):
    rows = [
        ("1", "Topic refinement and supervisor feedback", "Approved proposal direction"),
        ("2", "Literature review on BI, dashboards, RFM, clustering, and forecasting", "Literature notes"),
        ("3", "Dataset understanding and exploratory data analysis", "Dataset profile and data issues"),
        ("4", "Requirements analysis and dashboard module design", "Functional requirements and wireframe"),
        ("5", "Data cleaning and preprocessing pipeline", "Cleaned dataset"),
        ("6", "KPI calculator and executive overview module", "KPI outputs and overview page"),
        ("7", "Sales analytics module", "Sales charts and filters"),
        ("8", "Customer analytics and RFM calculation", "RFM table and customer metrics"),
        ("9", "Customer clustering implementation", "Segment labels and interpretation"),
        ("10", "Product analytics module", "Product performance views"),
        ("11", "Decision support insight logic", "Insight cards and alerts"),
        ("12", "Forecasting model implementation", "Forecast charts and metrics"),
        ("13", "Integration, UI improvement, and local testing", "Working prototype"),
        ("14", "Evaluation, validation, and deployment preparation", "Evaluation results and deployment notes"),
        ("15", "Documentation, README, and dissertation draft updates", "Technical documentation"),
        ("16", "Final testing, supervisor corrections, and presentation preparation", "Final demo and presentation materials"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["Week", "Planned Work", "Output"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        shade_cell(cell, "EDEDED")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_run_font(run, size=10, bold=True)

    for week, work, output in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, [week, work, output]):
            cell.text = text

    format_table(table, [Inches(0.6), Inches(3.7), Inches(2.2)])


def add_references(doc):
    references = [
        '[1] H. Chen, R. H. L. Chiang, and V. C. Storey, "Business Intelligence and Analytics: From Big Data to Big Impact," MIS Quarterly, vol. 36, no. 4, pp. 1165-1188, 2012.',
        "[2] S. Few, Information Dashboard Design: Displaying Data for At-a-Glance Monitoring, Analytics Press, 2006.",
        '[3] D. Chen, "Online Retail," UCI Machine Learning Repository, 2015. Available: https://archive.ics.uci.edu/dataset/352/online+retail',
        '[4] P. S. Fader, B. G. S. Hardie, and K. L. Lee, "RFM and CLV: Using Iso-Value Curves for Customer Base Analysis," Journal of Marketing Research, vol. 42, no. 4, pp. 415-430, 2005.',
        '[5] S. Lloyd, "Least Squares Quantization in PCM," IEEE Transactions on Information Theory, vol. 28, no. 2, pp. 129-137, 1982.',
        "[6] R. J. Hyndman and G. Athanasopoulos, Forecasting: Principles and Practice, 3rd ed., OTexts, 2021. Available: https://otexts.com/fpp3/",
        '[7] Scikit-learn Developers, "Scikit-learn: Machine Learning in Python." Available: https://scikit-learn.org/',
        '[8] Streamlit, "Streamlit Documentation." Available: https://docs.streamlit.io/',
    ]
    for ref in references:
        p = doc.add_paragraph()
        set_paragraph_format(p, after=6, line_spacing=1.15)
        run = p.add_run(ref)
        set_run_font(run, size=11)


def build_doc():
    doc = Document()
    setup_document(doc)

    add_para(doc, "FINAL YEAR INDIVIDUAL RESEARCH PROJECT PROPOSAL", after=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    title = doc.add_paragraph()
    set_paragraph_format(title, after=22, line_spacing=1.3, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run = title.add_run(TITLE)
    set_run_font(run, size=16, bold=True)

    metadata = [
        ("Student Name:", "[To be completed]"),
        ("Student ID:", "[To be completed]"),
        ("Degree Programme:", "[To be completed]"),
        ("Module Code:", "COM4901"),
        ("Supervisor:", "[To be completed]"),
        ("Faculty:", "Faculty of Computer Science and Engineering, KIU"),
        ("Submission Date:", "[To be completed]"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        set_paragraph_format(p, after=4, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        r1 = p.add_run(label + " ")
        set_run_font(r1, bold=True)
        r2 = p.add_run(value)
        set_run_font(r2)

    doc.add_page_break()

    add_heading(doc, "1. Introduction / Background")
    add_para(doc, "Small and Medium Enterprises (SMEs) often make strategic and operational decisions using sales records, customer information, product movement, and market performance. However, many smaller businesses do not have affordable or easy-to-use Business Intelligence (BI) systems that convert raw transaction data into useful management insights. As a result, important decisions such as identifying high-value customers, selecting profitable products, understanding sales trends, and planning future stock or marketing activities may be based on manual review or incomplete evidence.")
    add_para(doc, "Business Intelligence combines data collection, data cleaning, analytical processing, visualization, and reporting to support decision-making. In a retail context, a BI dashboard can help decision makers monitor Key Performance Indicators (KPIs), detect sales changes, identify product performance, and understand customer purchasing behavior. When BI is combined with simple analytics techniques such as customer segmentation and forecasting, the dashboard becomes a decision support tool rather than only a reporting tool.")
    add_para(doc, "This project proposes the design and development of a deployable web-based BI decision support dashboard for SME-style retail analysis. The project will use the publicly available UCI Online Retail dataset as a representative retail transaction dataset. The system will not use or invent private Sri Lankan SME data. Instead, it will demonstrate how an SME could use transaction records to support strategic decision-making while avoiding ethical and privacy risks associated with private business data.")
    add_para(doc, "The final system will be implemented as a Streamlit web application. It will include data cleaning, KPI calculation, interactive visual dashboards, customer analytics, product and sales analytics, decision-support insight cards, and at least one machine learning or analytics component such as customer segmentation and sales forecasting.")

    add_heading(doc, "2. Problem Statement")
    add_para(doc, "Many SMEs collect sales transaction data but lack an accessible decision support system that converts this data into clear, actionable insights. Existing spreadsheet-based analysis can be time-consuming, difficult to update, and limited in interactivity. Decision makers may struggle to answer questions such as:")
    add_bullets(doc, [
        "Which products generate the highest revenue?",
        "Which customers are most valuable or at risk of becoming inactive?",
        "Which markets or countries contribute most to sales?",
        "Are sales increasing or declining over time?",
        "What product or customer segments should the business prioritize?",
        "Can future sales trends be estimated using historical data?",
    ])
    add_para(doc, "Without a structured BI solution, these questions may be answered inconsistently or too late to support strategic decisions. Therefore, there is a need for a simple, low-cost, web-based BI dashboard that transforms retail transaction data into KPIs, visual analytics, customer segments, forecasts, and recommendation-style decision support.")

    add_heading(doc, "3. Research Aim and Objectives")
    add_heading(doc, "3.1 Aim", level=2)
    add_para(doc, "The aim of this project is to design, develop, and evaluate a Business Intelligence-based web dashboard that supports strategic decision-making for SME-style retail businesses using a publicly available transaction dataset.")
    add_heading(doc, "3.2 Objectives", level=2)
    add_numbers(doc, [
        "To investigate BI dashboard concepts, decision support systems, retail analytics, customer segmentation, and sales forecasting techniques relevant to SME decision-making.",
        "To preprocess and clean the UCI Online Retail dataset by handling missing values, cancelled or negative transactions, invalid prices, date features, and revenue calculations.",
        "To design and implement an interactive Streamlit dashboard with modules for executive overview, sales analytics, customer analytics, product analytics, decision support, and forecasting.",
        "To calculate business KPIs such as total revenue, total orders, total customers, average order value, product revenue, market contribution, and customer purchase frequency.",
        "To implement customer analytics using RFM analysis and clustering-based segmentation.",
        "To implement a simple forecasting or predictive analytics component using monthly aggregated revenue or quantity data, with model evaluation metrics.",
        "To evaluate the system using functional testing, data validation, usability-oriented review, dashboard usefulness criteria, and analytics/model performance metrics.",
    ])
    add_heading(doc, "3.3 Research Questions", level=2)
    add_numbers(doc, [
        "How can a low-cost BI dashboard be designed to support strategic decision-making for SME-style retail businesses?",
        "What KPIs and visual analytics are most useful for summarizing retail transaction performance?",
        "How can RFM analysis and clustering help identify meaningful customer segments?",
        "How accurately can simple forecasting models estimate short-term monthly sales trends from historical transaction data?",
        "How effectively does the developed dashboard convert raw transaction data into actionable decision-support insights?",
    ])

    add_heading(doc, "4. Scope and Deliverables")
    add_heading(doc, "4.1 Project Scope", level=2)
    add_para(doc, "The project scope is limited to the design and development of a prototype BI decision support dashboard using the UCI Online Retail dataset. The system will treat the dataset as a representative SME/retail transaction dataset and will focus on sales, customer, product, market, and forecasting insights.")
    add_bullets(doc, [
        "A Streamlit web application with multiple dashboard pages.",
        "Data loading, cleaning, preprocessing, and feature engineering.",
        "Executive KPIs and interactive charts.",
        "Sales analytics by month, country, product, customer, and quantity.",
        "Customer analytics using RFM and clustering.",
        "Product analytics for best-selling, slow-moving, and high-revenue products.",
        "Decision-support insight cards and simple recommendation logic.",
        "Forecasting or predictive analytics using monthly aggregated sales data.",
        "Local run instructions, GitHub-ready project structure, and Streamlit Cloud deployment documentation.",
    ])
    add_heading(doc, "4.2 Out of Scope", level=2)
    add_bullets(doc, [
        "Use of private Sri Lankan SME transaction data.",
        "Paid APIs, commercial BI tools, or enterprise data warehouses.",
        "Real-time integration with accounting, POS, ERP, or inventory systems.",
        "Advanced deep learning forecasting models.",
        "Production-level security, user account management, or multi-tenant business deployment.",
    ])
    add_heading(doc, "4.3 Expected Deliverables", level=2)
    add_bullets(doc, [
        "A functional Streamlit BI dashboard prototype.",
        "Cleaned and processed dataset files.",
        "Python source code organized into app.py, pages/, and src/ modules.",
        "Customer segmentation and forecasting outputs.",
        "Evaluation results including testing evidence and model performance metrics.",
        "README file with setup, usage, explanation, and deployment steps.",
        "Supporting documentation for methodology, dataset description, evaluation plan, and viva preparation.",
        "Final dissertation and presentation materials at the later project stage.",
    ])

    add_heading(doc, "5. Brief Literature Summary")
    add_para(doc, "Business Intelligence and analytics have become important for transforming organizational data into knowledge that supports decision-making. BI systems commonly include data integration, data analysis, performance measurement, dashboards, and reporting. For SMEs, BI can be valuable because it provides evidence-based insight without requiring large enterprise systems. However, BI adoption in smaller organizations can be limited by cost, technical complexity, and lack of analytical skills.")
    add_para(doc, "Dashboards are a common BI interface because they present performance information through KPIs, trends, comparisons, and alerts. A well-designed dashboard should help users understand business status quickly while allowing deeper exploration through filters and interactive charts. In a retail business, dashboard metrics such as revenue, order count, customer count, average order value, product revenue, and market contribution can support sales planning and product decisions.")
    add_para(doc, "Customer analytics is another important part of BI decision support. RFM analysis evaluates customers using recency, frequency, and monetary value. This method helps identify valuable customers, inactive customers, and customer groups requiring different marketing strategies. Clustering algorithms such as K-Means can be applied to RFM features to produce customer segments that support targeted decision-making.")
    add_para(doc, "Product and sales analytics support decisions about product focus, stock planning, market opportunities, and underperforming items. By comparing product revenue, quantity sold, market contribution, and demand patterns, businesses can identify products that should be promoted, monitored, or reviewed.")
    add_para(doc, "Forecasting supports decision-making by estimating future trends from historical data. Simple time series models such as moving averages, exponential smoothing, or regression-based approaches can be suitable for a student-level BI prototype because they are understandable, easy to evaluate, and suitable for monthly aggregated sales data. Forecasting performance can be assessed using metrics such as Mean Absolute Error (MAE), Root Mean Squared Error (RMSE), and Mean Absolute Percentage Error (MAPE), where applicable.")
    add_para(doc, "Based on this literature direction, the proposed project will combine BI dashboards, RFM-based customer segmentation, product analytics, and simple forecasting to create a practical decision support system for SME-style retail decision-making.")

    add_heading(doc, "6. Proposed Methodology / System Approach")
    add_para(doc, "This project will follow a development-oriented research approach. The methodology will include requirement analysis, dataset understanding, system design, implementation, testing, evaluation, and documentation.")
    add_heading(doc, "6.1 Dataset Understanding", level=2)
    add_para(doc, "The UCI Online Retail dataset contains retail transaction records from a UK-based online retail business. The attached Excel file contains 541,909 rows and 8 columns: InvoiceNo, StockCode, Description, Quantity, InvoiceDate, UnitPrice, CustomerID, and Country. The data covers transactions from 1 December 2010 to 9 December 2011. Initial inspection shows 38 countries, 4,372 identifiable customers, 25,900 invoices, and 4,070 stock codes. The dataset also contains cleaning issues such as missing customer IDs, negative quantity rows representing cancellations or returns, and zero-price rows.")
    add_heading(doc, "6.2 Data Cleaning and Preprocessing", level=2)
    add_bullets(doc, [
        "Remove or flag invalid records where quantity or unit price cannot support revenue analysis.",
        "Identify cancelled transactions using negative quantities and invoice numbers where applicable.",
        "Handle missing customer IDs depending on the analysis module.",
        "Standardize product descriptions and date fields.",
        "Create calculated features such as revenue, month, year-month, order value, and customer-level aggregates.",
        "Save cleaned outputs for repeatable dashboard use.",
    ])
    add_heading(doc, "6.3 System Design", level=2)
    add_para(doc, "The proposed system will be a modular Streamlit application. The main dashboard modules will be:")
    add_numbers(doc, [
        "Executive Overview: KPIs, monthly revenue trend, top countries, top products, and summary cards.",
        "Sales Analytics: sales trend, country performance, product performance, customer revenue, quantity analysis, and cancelled transaction handling.",
        "Customer Analytics: RFM analysis, customer segmentation, high-value customers, low-value customers, purchase frequency, and customer lifetime value-style indicators.",
        "Product Analytics: best-selling products, slow-moving products, high-revenue products, demand patterns, and product-level recommendations.",
        "Decision Support: business insight cards, alerts, and recommendation logic.",
        "Forecasting / Prediction: monthly revenue or quantity forecasting with model comparison and evaluation metrics.",
    ])
    add_heading(doc, "6.4 Analytics and Machine Learning", level=2)
    add_bullets(doc, [
        "Customer segmentation: RFM features will be calculated for each customer. StandardScaler will be used before clustering. K-Means clustering will be applied, and the number of clusters will be selected using simple evaluation support such as elbow method and silhouette score where feasible.",
        "Forecasting: monthly revenue or quantity will be aggregated and used to compare at least two simple forecasting models, such as moving average, linear regression trend, and exponential smoothing if feasible. Model performance will be measured using MAE, RMSE, and MAPE where applicable.",
    ])
    add_heading(doc, "6.5 Decision Support Logic", level=2)
    add_bullets(doc, [
        "If recent monthly revenue is lower than the previous period, show a declining sales alert.",
        "If a product has high revenue and high quantity sold, recommend prioritizing stock availability.",
        "If a product has low sales and low recent demand, recommend reviewing promotion or discontinuation.",
        "If a customer segment has high monetary value and high frequency, recommend retention campaigns.",
        "If a customer segment has high monetary value but low recent activity, recommend re-engagement.",
    ])
    add_heading(doc, "6.6 Evaluation Approach", level=2)
    add_bullets(doc, [
        "Functional testing: checking whether dashboard pages, filters, charts, and calculations work correctly.",
        "Data validation: comparing selected KPI outputs with manually calculated samples.",
        "Model evaluation: using MAE, RMSE, MAPE, silhouette score, and cluster interpretation where appropriate.",
        "Usability-oriented review: checking navigation, clarity of KPIs, chart readability, and usefulness of decision-support insights.",
        "Deployment verification: confirming that the app can run locally and can be prepared for Streamlit Cloud deployment.",
    ])

    add_heading(doc, "7. Tools and Resources Required")
    add_bullets(doc, [
        "Programming language: Python",
        "Web application framework: Streamlit",
        "Data processing: Pandas and NumPy",
        "Machine learning and evaluation: Scikit-learn",
        "Visualization: Plotly",
        "Dataset: UCI Online Retail dataset from the attached Online Retail.xlsx",
        "Development tools: VS Code or equivalent editor",
        "Version control: Git and GitHub",
        "Deployment target: Streamlit Cloud",
        "Documentation: Markdown files and final dissertation documents",
    ])

    add_heading(doc, "8. Ethical Considerations")
    add_para(doc, "This project will use a publicly available dataset for academic and development purposes. It will not collect new data from human participants and will not use private Sri Lankan SME data. The dataset contains customer identifiers, but these are numeric IDs and will be treated only as analytical labels for segmentation. No attempt will be made to identify real people or businesses. If any future supervisor-directed activity involves surveys, interviews, or real SME data collection, ethical approval and informed consent will be obtained before data collection.")
    add_para(doc, "The project will follow academic integrity requirements by citing dataset sources, literature, tools, and libraries. Code and documentation will be written by the student with assistance recorded where relevant. The project will avoid plagiarism and will maintain version control evidence.")

    add_heading(doc, "9. Project Plan / Timeline")
    add_para(doc, "The proposed project can be completed across 16 weeks. The timeline may be adjusted according to supervisor feedback and official university deadlines.")
    add_timeline_table(doc)

    add_heading(doc, "10. Expected Contribution")
    add_para(doc, "The expected contribution of this project is a practical, low-cost BI decision support dashboard prototype that demonstrates how retail transaction data can be transformed into actionable business insights for SME-style decision-making. The system will show the complete workflow from raw data cleaning to KPI reporting, customer segmentation, product analysis, forecasting, and recommendation-style decision support. Academically, the project will demonstrate the application of BI, data analytics, machine learning, and software development concepts in a complete final year project.")

    doc.add_page_break()
    add_heading(doc, "11. References")
    add_references(doc)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_doc()

