from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "output" / "Interim_Report_BI_Dashboard.docx"

TITLE = (
    "Design and Development of a Business Intelligence-Based Decision Support "
    "Dashboard for Strategic Decision-Making in Small and Medium Enterprises"
)

BLACK = RGBColor(0, 0, 0)
GRAY_FILL = "EDEDED"
LIGHT_FILL = "F7F7F7"
ACCENT_FILL = "FCE4D6"


def set_run_font(run, size=12, bold=False, italic=False, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def set_paragraph_format(paragraph, after=8, before=0, line_spacing=1.5, alignment=None):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing
    if alignment is not None:
        paragraph.alignment = alignment


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Page ")
    set_run_font(run, size=10)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def setup_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(8)

    for style_name, size, before, after in [
        ("Heading 1", 14, 14, 8),
        ("Heading 2", 13, 10, 6),
        ("Heading 3", 12, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLACK
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.5

    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def add_para(doc, text="", after=8, before=0, align=None, bold_label=None):
    p = doc.add_paragraph()
    set_paragraph_format(p, after=after, before=before, alignment=align)
    if bold_label:
        label = p.add_run(bold_label)
        set_run_font(label, bold=True)
        run = p.add_run(text)
        set_run_font(run)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size=14 if level == 1 else 13 if level == 2 else 12, bold=True)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_paragraph_format(p, after=4)
        run = p.add_run(item)
        set_run_font(run)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        set_paragraph_format(p, after=4)
        run = p.add_run(item)
        set_run_font(run)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def format_table(table, widths=None, header=True):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row_index, row in enumerate(table.rows):
        for idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if widths:
                cell.width = widths[idx]
            if header and row_index == 0:
                shade_cell(cell, GRAY_FILL)
            for paragraph in cell.paragraphs:
                set_paragraph_format(paragraph, after=0, line_spacing=1.15)
                for run in paragraph.runs:
                    set_run_font(run, size=10, bold=header and row_index == 0)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            cells[idx].text = str(value)
    format_table(table, widths)
    return table


def add_title_page(doc):
    add_para(doc, "INTERIM REPORT", after=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    p = doc.add_paragraph()
    set_paragraph_format(p, after=20, line_spacing=1.2, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(TITLE)
    set_run_font(run, size=16, bold=True)

    metadata = [
        ("Student Name:", "[To be completed]"),
        ("Student ID:", "[To be completed]"),
        ("Supervisor Name:", "[To be completed]"),
        ("Module Code:", "COM4901"),
        ("Faculty:", "Faculty of Computer Science and Engineering, KIU"),
        ("Submission Date:", "[To be completed]"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        set_paragraph_format(p, after=4, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        r1 = p.add_run(label + " ")
        set_run_font(r1, bold=True)
        r2 = p.add_run(value)
        set_run_font(r2)

    add_para(
        doc,
        "This interim report follows the COM4901 Final Year Individual Project guideline structure for interim progress reporting.",
        before=28,
        after=10,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    doc.add_page_break()


def add_introduction(doc):
    add_heading(doc, "1. Introduction")
    add_heading(doc, "1.1 Background and Context", level=2)
    add_para(
        doc,
        "Small and Medium Enterprises (SMEs) commonly generate sales transaction data through daily business activities. However, many SMEs do not have affordable and easy-to-use Business Intelligence (BI) tools that transform transaction records into reliable decision-support insights. In a retail setting, decisions about products, customers, markets, sales trends, and future planning can be improved when raw transaction data is converted into Key Performance Indicators (KPIs), visual analytics, customer segments, and forecasts.",
    )
    add_para(
        doc,
        "The proposed project focuses on the design and development of a BI-based decision support dashboard for SME-style retail analysis. The system will be implemented as a Streamlit web application using the attached UCI Online Retail dataset as a public representative business transaction dataset. The project does not use private Sri Lankan SME data and does not invent private business records.",
    )
    add_heading(doc, "1.2 Problem Statement", level=2)
    add_para(
        doc,
        "Many SMEs keep sales data but continue to rely on manual spreadsheet review or basic reports. This limits the ability to identify high-value customers, monitor product performance, detect market trends, and make strategic decisions based on evidence. Without an interactive decision-support dashboard, managers may not be able to answer important questions quickly, such as which products generate the most revenue, which customers should be retained, which markets perform best, and whether future sales may increase or decline.",
    )
    add_heading(doc, "1.3 Project Aim and Objectives", level=2)
    add_para(
        doc,
        "The aim of this project is to design, develop, and evaluate a BI-based web dashboard that supports strategic decision-making for SME-style retail businesses using a publicly available transaction dataset.",
    )
    add_numbers(
        doc,
        [
            "Investigate BI dashboard concepts, decision support systems, retail analytics, customer segmentation, and sales forecasting methods.",
            "Clean and preprocess the UCI Online Retail dataset for reliable analysis.",
            "Develop an interactive Streamlit dashboard with executive, sales, customer, product, decision-support, and forecasting modules.",
            "Calculate important business KPIs such as revenue, order count, customer count, average order value, market contribution, and product performance.",
            "Implement customer analytics using RFM analysis and clustering-based segmentation.",
            "Implement simple sales forecasting using monthly revenue or quantity data and evaluate model performance.",
            "Evaluate the dashboard using functional testing, data validation, model metrics, usability-oriented review, and deployment verification.",
        ],
    )


def add_progress_summary(doc):
    add_heading(doc, "2. Progress Summary")
    add_heading(doc, "2.1 Tasks Completed So Far", level=2)
    add_para(
        doc,
        "At the interim stage, the project has progressed through topic refinement, guideline review, proposal preparation, proposal presentation preparation, initial dataset understanding, literature direction planning, system module planning, and evaluation planning. The implementation stage is planned as the next major phase after proposal confirmation.",
    )
    rows = [
        ("Topic and scope definition", "Completed", "Project title, aim, objectives, scope, and deliverables drafted."),
        ("Guideline review", "Completed", "Proposal and interim report structure identified from COM4901 guidelines."),
        ("Dataset availability check", "Completed", "UCI Online Retail Excel dataset is available in the project workspace."),
        ("Initial dataset profiling", "Completed", "Rows, columns, dates, customer IDs, countries, invoices, and cleaning issues inspected."),
        ("Proposal document", "Completed", "Research proposal DOCX prepared for supervisor or panel approval."),
        ("Proposal presentation", "Completed", "PowerPoint proposal deck prepared for presentation."),
        ("System architecture planning", "In progress", "Planned Streamlit modules and source-code structure defined."),
        ("Prototype implementation", "Pending", "To begin after formal approval and supervisor feedback."),
    ]
    add_table(doc, ["Task", "Status", "Evidence / Output"], rows, [Inches(1.65), Inches(1.2), Inches(3.65)])

    add_heading(doc, "2.2 Current Project Status", level=2)
    add_para(
        doc,
        "The project is currently at the approved-concept and early design stage. The topic has been framed as a development-oriented project because it will produce a working web application prototype. The dataset has been inspected and confirmed suitable for business intelligence analysis. The next stage will focus on creating the project repository structure, implementing the preprocessing pipeline, calculating KPIs, and building the Streamlit dashboard modules.",
    )
    add_heading(doc, "2.3 Evidence of Progress", level=2)
    add_bullets(
        doc,
        [
            "Research proposal document prepared: output/Research_Proposal_BI_Dashboard.docx.",
            "Proposal presentation prepared: output/BI_Dashboard_Proposal_Presentation.pptx.",
            "UCI Online Retail dataset confirmed in the workspace: Online Retail.xlsx.",
            "Dataset profile recorded: 541,909 rows, 8 columns, 38 countries, 4,372 identifiable customer IDs, 25,900 invoice numbers, and 4,070 stock codes.",
            "Initial cleaning issues identified: 135,080 missing customer ID rows, 10,624 negative quantity rows, and 2,515 zero-price rows.",
        ],
    )


def add_literature_review_progress(doc):
    add_heading(doc, "3. Literature Review Progress")
    add_heading(doc, "3.1 Key Literature Identified", level=2)
    add_para(
        doc,
        "The literature review is currently focused on BI and analytics, dashboard design, customer analytics, clustering, and forecasting. These areas directly support the planned dashboard modules and provide the theoretical foundation for the system design.",
    )
    rows = [
        ("Business Intelligence and Analytics", "BI converts organizational data into decision-support knowledge through integration, analysis, visualization, and reporting.", "Supports the overall BI dashboard concept."),
        ("Dashboard Design", "Dashboards summarize performance using KPIs, trends, comparisons, and alerts for quick decision-making.", "Supports executive overview and interactive visual design."),
        ("RFM Analysis", "Recency, Frequency, and Monetary value are used to evaluate customer behavior and value.", "Supports customer segmentation and retention insights."),
        ("K-Means Clustering", "Clustering groups customers with similar RFM behavior into interpretable segments.", "Supports customer analytics and decision support."),
        ("Sales Forecasting", "Simple forecasting models estimate future trends from historical aggregated sales data.", "Supports the forecasting and prediction module."),
    ]
    add_table(doc, ["Area", "Main Idea", "Use in Project"], rows, [Inches(1.7), Inches(2.55), Inches(2.25)])

    add_heading(doc, "3.2 Theoretical and Conceptual Foundation", level=2)
    add_para(
        doc,
        "The conceptual foundation of the project is that raw transaction data can be transformed into decision-support value through a sequence of processing stages. First, the dataset is cleaned and prepared. Second, business KPIs are calculated. Third, dashboards and visualizations communicate patterns. Fourth, analytical models such as RFM segmentation and forecasting extend the system from reporting to decision support. Finally, rule-based insights translate the analytical outputs into practical recommendations.",
    )
    add_heading(doc, "3.3 Research Gap Justification", level=2)
    add_para(
        doc,
        "Enterprise BI platforms can be expensive and technically complex for smaller businesses. Many student and business prototypes also stop at visualization and do not include clear decision-support logic. This project addresses that gap by developing a low-cost, beginner-friendly BI dashboard prototype that combines KPIs, customer segmentation, product analytics, forecasting, and recommendation-style insight cards within a deployable Streamlit web application.",
    )


def add_methodology(doc):
    add_heading(doc, "4. Methodology / Solution Approach")
    add_heading(doc, "4.1 System Development Methodology", level=2)
    add_para(
        doc,
        "The project follows a development-oriented methodology. The major stages are requirements analysis, dataset understanding, system design, implementation, testing, evaluation, and documentation. This approach is suitable because the main output is a working dashboard prototype supported by academic explanation and evaluation.",
    )
    add_numbers(
        doc,
        [
            "Requirements analysis: identify business questions, dashboard pages, KPIs, and decision-support needs.",
            "Dataset understanding: inspect dataset structure, fields, date range, customers, countries, products, and data quality issues.",
            "System design: design modules for data loading, preprocessing, KPI calculation, visualization, segmentation, forecasting, and insights.",
            "Implementation: develop the Streamlit application and supporting Python modules.",
            "Testing and evaluation: validate calculations, check functionality, assess model results, and review usability.",
            "Documentation and deployment: prepare README, support documents, GitHub structure, and Streamlit Cloud deployment steps.",
        ],
    )
    add_heading(doc, "4.2 Data Collection / Dataset Approach", level=2)
    add_para(
        doc,
        "The project uses the public UCI Online Retail dataset supplied as Online Retail.xlsx. No new data collection is planned. The dataset contains invoice-level retail transactions with fields such as invoice number, stock code, product description, quantity, invoice date, unit price, customer ID, and country. It is suitable for sales analytics, product analytics, market analysis, customer segmentation, and monthly sales forecasting.",
    )
    add_heading(doc, "4.3 Tools, Techniques, and Justification", level=2)
    rows = [
        ("Python", "Core programming language", "Beginner-friendly and widely used for data analytics."),
        ("Pandas and NumPy", "Data cleaning and analysis", "Efficient handling of tabular transaction data."),
        ("Streamlit", "Web application framework", "Allows quick development of interactive dashboards."),
        ("Plotly", "Interactive visualizations", "Supports charts, filters, and presentation-quality visuals."),
        ("Scikit-learn", "Machine learning", "Supports K-Means clustering, scaling, and evaluation metrics."),
        ("GitHub", "Version control and source hosting", "Supports academic evidence and deployment preparation."),
        ("Streamlit Cloud", "Deployment", "Free deployment option suitable for a student project."),
    ]
    add_table(doc, ["Tool / Technique", "Purpose", "Justification"], rows, [Inches(1.55), Inches(2.05), Inches(2.9)])


def add_design_progress(doc):
    add_heading(doc, "5. Design and Implementation Progress")
    add_heading(doc, "5.1 Proposed System Architecture", level=2)
    add_para(
        doc,
        "The planned system architecture is modular. The raw Excel dataset will be loaded by a data loader module. A preprocessing module will clean the data and create calculated fields such as revenue and month. KPI, segmentation, forecasting, visualization, and decision-support modules will use the cleaned data to generate dashboard outputs. Streamlit pages will present these outputs through an interactive web interface.",
    )
    rows = [
        ("Layer 1", "Data Source", "Online Retail.xlsx public transaction dataset."),
        ("Layer 2", "Data Processing", "Load, clean, validate, and engineer features."),
        ("Layer 3", "Analytics", "KPIs, RFM, clustering, product analysis, and forecasting."),
        ("Layer 4", "Decision Support", "Insight cards, alerts, and recommendation rules."),
        ("Layer 5", "User Interface", "Streamlit dashboard pages and interactive charts."),
    ]
    add_table(doc, ["Layer", "Component", "Description"], rows, [Inches(0.9), Inches(1.75), Inches(3.85)])

    add_heading(doc, "5.2 Planned Repository Structure", level=2)
    add_para(
        doc,
        "The application will be organized so that the main Streamlit application is easy to run and the analytics logic is separated into reusable source modules.",
    )
    rows = [
        ("app.py", "Main Streamlit entry point."),
        ("pages/", "Dashboard pages such as executive overview, sales analytics, customer analytics, product analytics, decision support, and forecasting."),
        ("src/data_loader.py", "Dataset loading functions."),
        ("src/preprocessing.py", "Cleaning and feature engineering functions."),
        ("src/kpi_calculator.py", "Business KPI calculations."),
        ("src/segmentation.py", "RFM and clustering logic."),
        ("src/forecasting.py", "Monthly forecasting models and metrics."),
        ("src/visualizations.py", "Reusable Plotly chart functions."),
        ("data/raw/", "Original dataset storage."),
        ("data/processed/", "Cleaned and prepared data outputs."),
        ("models/", "Saved model or clustering outputs if needed."),
        ("README.md", "Setup, usage, explanation, and deployment instructions."),
    ]
    add_table(doc, ["File / Folder", "Purpose"], rows, [Inches(2.15), Inches(4.35)])

    add_heading(doc, "5.3 Implementation Progress", level=2)
    add_para(
        doc,
        "At this stage, full dashboard implementation has not yet been completed because the project is being prepared for formal approval before development continues. Completed technical preparation includes dataset availability checking, initial data profiling, requirement definition, dashboard module planning, analytics planning, and creation of proposal artefacts. The next implementation milestone is to create the repository structure and begin the preprocessing pipeline.",
    )
    add_heading(doc, "5.4 Screenshots / Outputs", level=2)
    add_para(
        doc,
        "Screenshots of the working Streamlit dashboard will be inserted after implementation. Planned screenshot evidence includes the home page, executive overview, sales analytics page, customer segmentation page, product analytics page, decision-support page, forecasting page, and deployment page or local run evidence.",
    )


def add_evaluation_plan(doc):
    add_heading(doc, "6. Testing / Evaluation Plan")
    add_heading(doc, "6.1 Proposed Testing Strategy", level=2)
    rows = [
        ("Unit testing / function checks", "Check preprocessing, KPI, RFM, clustering, and forecasting functions.", "Correct outputs for known sample inputs."),
        ("Data validation", "Compare selected dashboard KPIs with manual Pandas calculations.", "Revenue, order count, customer count, and AOV match expected values."),
        ("Functional testing", "Check dashboard navigation, filters, charts, and page loading.", "All pages load and interactive controls work."),
        ("Model evaluation", "Assess customer clustering and forecasting models.", "Silhouette score and forecast error metrics reported where applicable."),
        ("Usability-oriented review", "Review readability, chart clarity, and decision-support usefulness.", "Dashboard is clear enough for SME-style decision-making."),
        ("Deployment verification", "Run locally and prepare for Streamlit Cloud.", "App starts successfully and requirements are documented."),
    ]
    add_table(doc, ["Test Type", "Description", "Success Evidence"], rows, [Inches(1.5), Inches(2.55), Inches(2.45)])

    add_heading(doc, "6.2 Evaluation Metrics", level=2)
    add_bullets(
        doc,
        [
            "Dashboard KPIs: total revenue, total orders, total customers, average order value, product revenue, customer revenue, and country revenue.",
            "Customer segmentation: RFM feature interpretation, cluster profiles, and silhouette score where feasible.",
            "Forecasting: Mean Absolute Error (MAE), Root Mean Squared Error (RMSE), and Mean Absolute Percentage Error (MAPE) where suitable.",
            "Functional quality: number of working pages, successful file load, successful filter interactions, and chart rendering status.",
            "Deployment readiness: requirements.txt, README instructions, GitHub structure, and Streamlit Cloud deployment steps.",
        ],
    )
    add_heading(doc, "6.3 Planned Experiments or Performance Measures", level=2)
    add_para(
        doc,
        "For forecasting, monthly revenue or quantity will be aggregated from the cleaned dataset. At least two simple models will be compared, such as moving average and linear regression trend. If feasible, exponential smoothing will also be tested. The final model will not be selected only by accuracy; interpretability and suitability for a student-level BI dashboard will also be considered.",
    )


def add_challenges_and_risks(doc):
    add_heading(doc, "7. Challenges and Risk Management")
    rows = [
        ("Large dataset size", "The Excel file contains more than 500,000 transaction rows.", "Use cached loading, processed CSV files, and efficient Pandas operations."),
        ("Missing customer IDs", "Some rows cannot be used for customer-level RFM analysis.", "Use only valid customer IDs for segmentation and document excluded rows."),
        ("Cancelled or negative transactions", "Negative quantities can distort revenue analysis.", "Separate, flag, or exclude returns depending on the analysis page."),
        ("Forecasting limitations", "Only about one year of monthly data is available.", "Use simple baseline models and clearly explain limitations."),
        ("Scope creep", "Too many dashboard features may reduce completion quality.", "Prioritize required modules and keep implementation beginner-friendly."),
        ("Deployment issues", "Streamlit Cloud may fail if file paths or dependencies are not prepared correctly.", "Use requirements.txt, relative paths, and deployment testing."),
    ]
    add_table(doc, ["Challenge / Risk", "Explanation", "Mitigation"], rows, [Inches(1.65), Inches(2.35), Inches(2.5)])
    add_para(
        doc,
        "The main corrective action is to keep the project modular and test each module separately. This reduces the risk of late-stage integration issues and allows progress evidence to be demonstrated even if some advanced features need to be simplified.",
    )


def add_revised_work_plan(doc):
    add_heading(doc, "8. Revised Work Plan")
    add_para(
        doc,
        "The revised work plan prioritizes implementation and evaluation after approval. The schedule may be adjusted according to supervisor feedback and official faculty deadlines.",
    )
    rows = [
        ("1", "Confirm proposal approval and supervisor corrections", "Approved topic and refined scope"),
        ("2", "Create repository structure and requirements file", "GitHub-ready project skeleton"),
        ("3", "Build data loader and preprocessing module", "Cleaned processed dataset"),
        ("4", "Implement KPI calculations and executive overview", "Core KPI dashboard page"),
        ("5", "Implement sales analytics and product analytics", "Sales and product pages"),
        ("6", "Implement RFM analysis and customer segmentation", "Customer analytics page"),
        ("7", "Implement decision-support rules and insight cards", "Decision support page"),
        ("8", "Implement forecasting and model evaluation", "Forecasting page and metrics"),
        ("9", "Integrate pages, improve UI, and test locally", "Working Streamlit prototype"),
        ("10", "Prepare README, deployment instructions, and report evidence", "Documentation and final evidence"),
    ]
    add_table(doc, ["Stage", "Remaining Task", "Expected Output"], rows, [Inches(0.75), Inches(3.35), Inches(2.4)])


def add_conclusion(doc):
    add_heading(doc, "9. Conclusion")
    add_para(
        doc,
        "The project remains feasible and aligned with the COM4901 development-oriented project requirements. The proposal, presentation, dataset understanding, literature direction, methodology, and system design have been prepared. The next phase will focus on implementation of the Streamlit dashboard, analytics modules, testing, evaluation, and deployment preparation.",
    )
    add_para(
        doc,
        "The use of a public dataset avoids ethical risks linked to private SME data while still allowing realistic retail business analysis. The planned system is technically achievable using Python, Pandas, Scikit-learn, Plotly, and Streamlit. With careful scope control and continuous testing, the project can progress from proposal approval to a complete deployable BI decision support dashboard.",
    )


def add_references(doc):
    doc.add_page_break()
    add_heading(doc, "10. References")
    refs = [
        '[1] H. Chen, R. H. L. Chiang, and V. C. Storey, "Business Intelligence and Analytics: From Big Data to Big Impact," MIS Quarterly, vol. 36, no. 4, pp. 1165-1188, 2012.',
        "[2] S. Few, Information Dashboard Design: Displaying Data for At-a-Glance Monitoring, Analytics Press, 2006.",
        '[3] D. Chen, "Online Retail," UCI Machine Learning Repository, 2015. Available: https://archive.ics.uci.edu/dataset/352/online+retail',
        '[4] P. S. Fader, B. G. S. Hardie, and K. L. Lee, "RFM and CLV: Using Iso-Value Curves for Customer Base Analysis," Journal of Marketing Research, vol. 42, no. 4, pp. 415-430, 2005.',
        '[5] S. Lloyd, "Least Squares Quantization in PCM," IEEE Transactions on Information Theory, vol. 28, no. 2, pp. 129-137, 1982.',
        "[6] R. J. Hyndman and G. Athanasopoulos, Forecasting: Principles and Practice, 3rd ed., OTexts, 2021. Available: https://otexts.com/fpp3/",
        '[7] Scikit-learn Developers, "Scikit-learn: Machine Learning in Python." Available: https://scikit-learn.org/',
        '[8] Streamlit, "Streamlit Documentation." Available: https://docs.streamlit.io/',
    ]
    for ref in refs:
        p = doc.add_paragraph()
        set_paragraph_format(p, after=6, line_spacing=1.15)
        run = p.add_run(ref)
        set_run_font(run, size=11)


def add_appendices(doc):
    doc.add_page_break()
    add_heading(doc, "11. Appendices")
    add_heading(doc, "Appendix A: Dataset Field Description", level=2)
    rows = [
        ("InvoiceNo", "Invoice number for each transaction."),
        ("StockCode", "Product code."),
        ("Description", "Product description."),
        ("Quantity", "Number of items purchased or returned."),
        ("InvoiceDate", "Transaction date and time."),
        ("UnitPrice", "Price per unit."),
        ("CustomerID", "Numeric customer identifier."),
        ("Country", "Customer country or market."),
    ]
    add_table(doc, ["Field", "Description"], rows, [Inches(1.6), Inches(4.9)])

    add_heading(doc, "Appendix B: Planned Screenshot Evidence", level=2)
    add_bullets(
        doc,
        [
            "[Insert screenshot] Streamlit home or executive overview page.",
            "[Insert screenshot] Sales analytics page with monthly trend.",
            "[Insert screenshot] Customer analytics page with RFM segmentation.",
            "[Insert screenshot] Product analytics page.",
            "[Insert screenshot] Decision support insight cards.",
            "[Insert screenshot] Forecasting page with model comparison and metrics.",
        ],
    )

    add_heading(doc, "Appendix C: Project Diary / Meeting Evidence Placeholder", level=2)
    rows = [
        ("[Date]", "[Meeting or work session]", "[Supervisor feedback / completed task]", "[Next action]"),
        ("[Date]", "[Meeting or work session]", "[Supervisor feedback / completed task]", "[Next action]"),
        ("[Date]", "[Meeting or work session]", "[Supervisor feedback / completed task]", "[Next action]"),
    ]
    add_table(doc, ["Date", "Activity", "Evidence / Feedback", "Action Point"], rows, [Inches(1.1), Inches(1.6), Inches(2.3), Inches(1.5)])


def build_doc():
    doc = Document()
    setup_document(doc)
    add_title_page(doc)
    add_introduction(doc)
    add_progress_summary(doc)
    add_literature_review_progress(doc)
    add_methodology(doc)
    add_design_progress(doc)
    add_evaluation_plan(doc)
    add_challenges_and_risks(doc)
    add_revised_work_plan(doc)
    add_conclusion(doc)
    add_references(doc)
    add_appendices(doc)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_doc()
